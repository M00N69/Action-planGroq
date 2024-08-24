import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from fpdf import FPDF
from groq import Groq

# Configuration de la page
st.set_page_config(layout="wide")

# Fonction pour ajouter des styles CSS personnalisés
def add_css_styles():
    st.markdown(
        """
        <style>
        .main-header {
            font-size: 28px;
            font-weight: bold;
            color: #2c3e50;
            text-align: center;
            margin-bottom: 30px;
        }
        .banner {
            text-align: center;
            margin-bottom: 40px;
        }
        .dataframe-container {
            margin: 20px 0;
            border: 1px solid #ddd;
            padding: 10px;
            border-radius: 10px;
            background-color: #f9f9f9;
        }
        .button-column {
            text-align: right;
        }
        .recommendation-container {
            padding: 15px;
            border: 1px solid #3498db;
            border-radius: 5px;
            background-color: #ecf0f1;
            margin-top: 20px;
        }
        .recommendation-header {
            font-weight: bold;
            font-size: 18px;
            color: #2980b9;
            margin-bottom: 10px;
        }
        .recommendation-content {
            margin-bottom: 10px;
            font-size: 16px;
            line-height: 1.5;
        }
        .expander {
            margin-top: 10px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

# Fonction pour configurer le client Groq
def get_groq_client():
    """Initialise et renvoie un client Groq avec la clé API."""
    return Groq(api_key=st.secrets["GROQ_API_KEY"])

# Fonction pour charger le fichier Excel avec le plan d'action
def load_action_plan(uploaded_file):
    try:
        action_plan_df = pd.read_excel(uploaded_file, header=12)
        expected_columns = ["Numéro d'exigence", "Exigence IFS Food 8", "Notation", "Explication (par l’auditeur/l’évaluateur)"]
        action_plan_df = action_plan_df[expected_columns]
        return action_plan_df
    except Exception as e:
        st.error(f"Erreur lors de la lecture du fichier: {str(e)}")
        return None

# Fonction pour générer un prompt basé sur une non-conformité
def generate_ai_recommendation_groq(non_conformity, guide_row):
    client = get_groq_client()
    prompt = f"""
    En tant qu'expert en IFS Food 8, fournissez des recommandations pour :
    - Exigence : {non_conformity["Numéro d'exigence"]}
    - Description : {non_conformity['Exigence IFS Food 8']}
    - Constat détaillé : {non_conformity['Explication (par l’auditeur/l’évaluateur)']}
    - Bonnes pratiques : {guide_row['Good practice']}
    - Éléments à vérifier : {guide_row['Elements to check']}
    - Exemple de question : {guide_row['Example questions']}
    """
    
    messages = [
        {"role": "system", "content": "Générer une recommandation structurée pour la non-conformité suivante :"},
        {"role": "user", "content": prompt}
    ]

    try:
        chat_completion = client.chat.completions.create(
            messages=messages,
            model="llama-3.1-70b-versatile"
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        st.error(f"Erreur lors de la génération de la recommandation: {str(e)}")
        return None

# Fonction pour récupérer les informations du guide
def get_guide_info(num_exigence, guide_df):
    return guide_df[guide_df['NUM_REQ'] == num_exigence].iloc[0]

# Fonction pour afficher les recommandations
def display_recommendations(recommendations_df):
    for index, row in recommendations_df.iterrows():
        with st.expander(f"Voir la recommandation pour l'exigence {row['Numéro d'exigence']}"):
            st.markdown(f"""**Numéro d'exigence:** {row["Numéro d'exigence"]}""")
            st.markdown(row["Recommandation"])

# Fonction pour créer un fichier texte des recommandations
def generate_text_file(recommendations_df):
    text_content = ""
    for index, row in recommendations_df.iterrows():
        text_content += f"""Numéro d'exigence: {row["Numéro d'exigence"]}\n"""
        text_content += f"""{row["Recommandation"]}\n\n"""
    return text_content

# Fonction pour créer un fichier DOCX des recommandations
def generate_docx_file(recommendations_df):
    doc = Document()
    for index, row in recommendations_df.iterrows():
        doc.add_heading(f"""Numéro d'exigence: {row["Numéro d'exigence"]}""", level=2)
        doc.add_paragraph(row["Recommandation"])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Fonction pour créer un fichier PDF des recommandations
def generate_pdf_file(recommendations_df):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for index, row in recommendations_df.iterrows():
        pdf.set_font("Arial", style='B', size=14)
        pdf.cell(200, 10, txt=f"""Numéro d'exigence: {row["Numéro d'exigence"]}""", ln=True)
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=row["Recommandation"])
        pdf.ln(10)
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer

# Fonction principale
def main():
    add_css_styles()
    
    st.markdown('<div class="main-header">Assistant VisiPilot pour Plan d\'Actions IFS</div>', unsafe_allow_html=True)
    
    # Ajout de la section d'explication
    with st.expander("Comment utiliser cette application"):
        st.write("""
            **Étapes d'utilisation:**
            1. Téléchargez votre plan d'action IFSv8.
            2. Sélectionnez un numéro d'exigence.
            3. Cliquez sur "Générer Recommandations" pour obtenir des suggestions.
            4. Téléchargez les recommandations générées.
        """)

    st.write("Téléchargez votre plan d'action et obtenez des recommandations pour les corrections et actions correctives.")

    # Upload du fichier Excel
    uploaded_file = st.file_uploader("Téléchargez votre plan d'action (fichier Excel)", type=["xlsx"])
    
    if uploaded_file:
        action_plan_df = load_action_plan(uploaded_file)
        if action_plan_df is not None:
            st.markdown('<div class="dataframe-container">' + action_plan_df.to_html(classes='dataframe', index=False) + '</div>', unsafe_allow_html=True)
            
            # Charger le guide IFSv8 depuis un fichier CSV
            guide_df = pd.read_csv("https://raw.githubusercontent.com/M00N69/Action-planGroq/main/Guide%20Checklist_IFS%20Food%20V%208%20-%20CHECKLIST.csv")

            recommendations = []
            selected_exigence = st.selectbox("Sélectionnez un numéro d'exigence :", action_plan_df["Numéro d'exigence"].unique())

            button_col, _ = st.columns([1, 2])
            with button_col:
                if st.button("Générer Recommandations"):
                    guide_row = get_guide_info(selected_exigence, guide_df)
                    non_conformity = action_plan_df[action_plan_df["Numéro d'exigence"] == selected_exigence].iloc[0]
                    
                    with st.spinner('Génération des recommandations...'):
                        recommendation_text = generate_ai_recommendation_groq(non_conformity, guide_row)
                        if recommendation_text:
                            st.success("Recommandation générée avec succès!")
                            recommendations.append({
                                "Numéro d'exigence": selected_exigence,
                                "Recommandation": recommendation_text
                            })
                            action_plan_df.loc[action_plan_df["Numéro d'exigence"] == selected_exigence, "Recommandation"] = recommendation_text

            if recommendations:
                recommendations_df = pd.DataFrame(recommendations)
                st.subheader("Résumé des Recommandations")
                display_recommendations(recommendations_df)

                # Options de téléchargement
                st.download_button(
                    label="Télécharger les recommandations (CSV)",
                    data=recommendations_df.to_csv(index=False),
                    file_name="recommandations_ifs_food.csv",
                    mime="text/csv"
                )

                text_file = generate_text_file(recommendations_df)
                st.download_button(
                    label="Télécharger les recommandations (Texte)",
                    data=text_file,
                    file_name="recommandations_ifs_food.txt",
                    mime="text/plain"
                )

                docx_file = generate_docx_file(recommendations_df)
                st.download_button(
                    label="Télécharger les recommandations (DOCX)",
                    data=docx_file,
                    file_name="recommandations_ifs_food.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                pdf_file = generate_pdf_file(recommendations_df)
                st.download_button(
                    label="Télécharger les recommandations (PDF)",
                    data=pdf_file,
                    file_name="recommandations_ifs_food.pdf",
                    mime="application/pdf"
                )

if __name__ == "__main__":
    main()
